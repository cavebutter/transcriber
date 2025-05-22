import requests
import json
import re
import os
import docx
import argparse
import jinja2
import datetime
import shutil
from md_to_pdf import convert_md_to_pdf

address = 'http://localhost:11434/api/generate'
model = None
payload = {
    'model': model,
    'prompt': None,
    'stream': False,
}

bullet_prompt = """What follows is a diarized transcript of a meeting between a number of people in the electric utility
industry in the United States. Topics discussed will be related to that or related to project management and software
development activities to support the electric utility. 

Your output should be in proper markdown format with the following:
1. Use ** for bold text (e.g., **important point**)
2. Use ## for topic headings (level 2 headers) - DO NOT use numbered sections or ### headers
3. Use - for bulleted lists (always use a dash followed by a space)
4. Use blank lines between paragraphs
5. Never use numbered lists for bullet points
6. DO NOT include any numbering (1.0.1, 2.1, etc.) in headers or sections
7. DO NOT create subsections under topic headers

You must include the following sections: 
1. For each key topic discussed: a descriptive name as a level 2 header (##), followed by a bulleted list of key points using dashes (-), with action items highlighted in bold.

Do not include anything else except for the above. No extraneous words, HTML formatting, or numbered sections. \n \n"""



exec_summary_prompt = """What follows is a list of key topics discussed in a meeting between a number of people in the 
electric utility industry. The topic will be related to this, or to project management and software development 
activities to support the electric utility.

Your output MUST BE EXACTLY ONE PARAGRAPH that summarizes the meeting.
- Use plain text with markdown formatting (use ** for bold, not HTML tags)
- Write as a single continuous paragraph with no line breaks
- Focus only on the most critical information
- Keep your summary concise (under 150 words)
- Do not include bullet points, numbered lists, or multiple paragraphs
- No extraneous words or HTML formatting at all

Write only the executive summary paragraph and nothing else.\n \n"""


# Keep the existing prompts for generating the content in markdown format
# We'll convert the markdown to HTML later

def parse_args():
    """Parse command-line arguments."""
    parser = argparse.ArgumentParser(description="Summarize a meeting transcript.")
    parser.add_argument('transcript_file', type=str, help='Path to the transcript file (.txt or .docx)')
    parser.add_argument('--title', '-t', type=str, help='Title of the meeting')
    parser.add_argument('--date', '-d', type=str, help='Date of the meeting')
    parser.add_argument('--output-dir', '-o', type=str, default='.', help='Directory to save the output files')
    parser.add_argument('--format', '-f', type=str, choices=['md', 'html', 'pdf'], default='html',
                        help='Output format (md=markdown, html=HTML, pdf=PDF)')
    parser.add_argument('--model', '-m', type=str, default='qwen3-summarizer:14b',
                        choices=['qwen3-summarizer:14b', 'qwen3-summarizer:30b'],
                        help='Model to use for summarization')
    return parser.parse_args()

def load_diarized_transcript(transcript_file):
    """
    Load a transcript file based on its extension (.txt or .docx)

    Returns:
        tuple: (transcript_text, participants) where participants is a list or None
    """
    _, ext = os.path.splitext(transcript_file)
    participants = None

    if ext.lower() == '.txt':
        with open(transcript_file) as f:
            transcript = f.read()
    elif ext.lower() == '.docx':
        transcript = extract_text_from_docx(transcript_file)
        participants = extract_participants_from_teams_docx(transcript_file)
    else:
        raise ValueError(f"Unsupported file format: {ext}. Only .txt and .docx are supported.")

    return transcript, participants


def extract_text_from_docx(docx_file):
    """
    Extract text content from MS Teams transcript (.docx) file,
    ignoring images and preserving speaker information and timestamps.
    """
    print(f"Loading MS Teams transcript from {docx_file}...")
    doc = docx.Document(docx_file)
    transcript_text = []

    # Process paragraphs to extract speaker info and text
    current_speaker = None

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # In Teams transcripts, speaker is typically in bold text
        # followed by timestamp and then the actual content
        if ':' in text and len(text.split(':')[0]) < 50:  # Simple heuristic for speaker line
            parts = text.split(':', 1)
            speaker_info = parts[0].strip()
            content = parts[1].strip() if len(parts) > 1 else ""

            # Extract speaker name (might contain timestamp too)
            current_speaker = speaker_info
            if content:
                transcript_text.append(f"{current_speaker}: {content}")
        else:
            # This is content from the previous speaker
            if current_speaker:
                transcript_text.append(f"{current_speaker}: {text}")
            else:
                transcript_text.append(text)

    return '\n'.join(transcript_text)


def extract_participants_from_teams_docx(docx_file):
    """
    Extract a list of unique participants from MS Teams transcript (.docx) file.
    Cleans up names to remove timestamps and numeric indicators.

    Args:
        docx_file (str): Path to the MS Teams transcript docx file

    Returns:
        list: A list of unique participant names
    """
    print(f"Extracting participants from MS Teams transcript: {docx_file}...")
    doc = docx.Document(docx_file)
    participants = set()

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # In Teams transcripts, speaker is typically before the colon
        if ':' in text and len(text.split(':')[0]) < 50:  # Simple heuristic for speaker line
            speaker_info = text.split(':', 1)[0].strip()

            # Clean up speaker info by extracting just the name
            # Remove timestamps like (10:45) and numeric indicators like "35"
            speaker_name = speaker_info

            # Remove content in parentheses (like timestamps)
            if '(' in speaker_name:
                speaker_name = speaker_name.split('(')[0].strip()

            # Remove any trailing numbers after the name with spaces
            speaker_name = re.sub(r'\s+\d+\s*$', '', speaker_name)

            # Only add non-empty names and avoid date-like entries
            if speaker_name and not re.match(
                    r'^(January|February|March|April|May|June|July|August|September|October|November|December)',
                    speaker_name):
                participants.add(speaker_name)

    # Convert set to sorted list
    return sorted(list(participants))


def ollama_request_think_tags(payload_: dict):
    response = requests.post(address, json=payload)
    if response.status_code == 200:
        response_data = response.json()
        full_response = response_data['response']

        match = re.search(r'</think>(.*)', full_response, re.DOTALL)
        if match:
            text_after_think = match.group(1).strip()
            return text_after_think
        else:
            return full_response
    else:
        return None


def get_bullet_summary(transcript_file: str, model_name: str):
    transcript, participants = load_diarized_transcript(transcript_file)

    # Update prompt with participant information if available
    local_bullet_prompt = bullet_prompt
    if participants:
        participants_str = ", ".join(participants)
        local_bullet_prompt += f"Meeting participants: {participants_str}\n\n"

    payload['model'] = model_name
    payload['prompt'] = local_bullet_prompt + transcript
    bullet_summary = ollama_request_think_tags(payload)
    return bullet_summary, participants


def get_exec_summary(bullet_summary: str, model_name: str):
    payload['model'] = model_name
    payload['prompt'] = exec_summary_prompt + bullet_summary
    exec_summary = ollama_request_think_tags(payload)
    return exec_summary


def get_markdown_document(exec_summary: str, bullet_summary: str, output_file: str, participants=None):
    """Generate a markdown document and save to file"""
    with open(output_file, 'w') as f:
        # Executive Summary Section
        f.write("# Executive Summary\n\n")

        # Clean and ensure single paragraph
        cleaned_exec = exec_summary.strip()
        # Remove any headers or numbering that might have been added
        cleaned_exec = re.sub(r'^#+\s*.*$', '', cleaned_exec, flags=re.MULTILINE)
        cleaned_exec = re.sub(r'^\d+\.[\d\.]*\s*.*$', '', cleaned_exec, flags=re.MULTILINE)
        cleaned_exec = re.sub(r'\n\s*\n', ' ', cleaned_exec)  # Collapse multiple newlines
        cleaned_exec = cleaned_exec.strip()

        f.write(cleaned_exec)
        f.write('\n\n')

        # Participants Section if available
        if participants and len(participants) > 0:
            f.write("## Meeting Participants\n\n")
            # Create 4-column table for participants without headers
            cols = 4
            rows = (len(participants) + cols - 1) // cols

            f.write("| | | | |\n")
            f.write("| --- | --- | --- | --- |\n")

            for i in range(rows):
                row = "| "
                for j in range(cols):
                    idx = i + j * rows
                    if idx < len(participants):
                        row += participants[idx] + " | "
                    else:
                        row += " | "
                f.write(row.strip() + "\n")
            f.write('\n\n')

        # Key Discussion Topics Section
        f.write("# Key Discussion Topics and Decisions\n\n")

        # Clean up the bullet summary
        cleaned_summary = bullet_summary

        # Remove HTML tags
        cleaned_summary = cleaned_summary.replace("<strong>", "**").replace("</strong>", "**")
        cleaned_summary = cleaned_summary.replace("<ul>", "").replace("</ul>", "")
        cleaned_summary = cleaned_summary.replace("<li>", "- ").replace("</li>", "")
        cleaned_summary = cleaned_summary.replace("<p>", "").replace("</p>", "\n\n")

        # Remove numbered sections and convert to proper headers
        cleaned_summary = re.sub(r'^\d+\.[\d\.]*\s*(.*?)$', r'## \1', cleaned_summary, flags=re.MULTILINE)

        # Ensure headers use ## not ###
        cleaned_summary = cleaned_summary.replace("### ", "## ")

        f.write(cleaned_summary)

    return exec_summary, bullet_summary, participants

def markdown_to_html(title, date, exec_summary, bullet_summary, participants, output_dir="."):
    """Convert markdown content to HTML using Jinja2"""
    # First, ensure bullet_summary is a string
    if isinstance(bullet_summary, list):
        bullet_summary = "\n".join(bullet_summary)

    # Define template and CSS locations
    current_dir = os.path.dirname(os.path.abspath(__file__))
    template_dir = os.path.join(current_dir, "templates")
    template_name = "meeting_template.html"
    css_name = "template.css"

    # Create templates directory if it doesn't exist
    if not os.path.exists(template_dir):
        os.makedirs(template_dir)

    template_path = os.path.join(template_dir, template_name)
    css_path = os.path.join(template_dir, css_name)

    # If template doesn't exist, create it
    if not os.path.exists(template_path):
        create_html_template(template_path)

    # If CSS doesn't exist, create it
    if not os.path.exists(css_path):
        create_css_template(css_path)

    # Copy CSS to output directory - this is the key change
    output_css_path = os.path.join(output_dir, css_name)
    shutil.copy2(css_path, output_css_path)
    print(f"Copied CSS file to: {output_css_path}")

    # Setup Jinja2 environment
    template_loader = jinja2.FileSystemLoader(searchpath=template_dir)
    template_env = jinja2.Environment(loader=template_loader)

    # Load template
    template = template_env.get_template(template_name)

    # Convert bullet summary to HTML
    bullet_html = markdown_to_html_content(bullet_summary)

    # Format participants into a table with up to 4 columns, no headers
    participants_html = ""
    if participants:
        participants_html = "<table class='participants-table'>"
        cols = 4
        rows = (len(participants) + cols - 1) // cols  # Ceiling division

        for i in range(rows):
            participants_html += "<tr>"
            for j in range(cols):
                idx = i + j * rows
                if idx < len(participants):
                    participants_html += f"<td>{participants[idx]}</td>"
                else:
                    participants_html += "<td></td>"
            participants_html += "</tr>"
        participants_html += "</table>"

    # Get current date if not provided
    if not date:
        date = datetime.datetime.now().strftime("%B %d, %Y")

    # Render template - the HTML should use a relative path to CSS
    html_content = template.render(
        title=title or "Meeting Summary",
        date=date,
        executive_summary=exec_summary,
        participants_html=participants_html,
        content=bullet_html,
        current_year=datetime.datetime.now().year
    )

    return html_content


if __name__ == "__main__":
    args = parse_args()

    # Get transcript file from args
    transcript_file = args.transcript_file
    model_name = args.model

    # Create base output filename with same name but different extension in the specified output directory
    base_name = os.path.basename(os.path.splitext(transcript_file)[0])

    # Make sure output directory exists
    if not os.path.exists(args.output_dir):
        os.makedirs(args.output_dir)

    print(f"Processing transcript: {transcript_file} with model {model_name}")

    # Process the transcript
    bullet_summary, participants = get_bullet_summary(transcript_file, model_name)
    exec_summary = get_exec_summary(bullet_summary, model_name)

    # Output handling based on selected format
    if args.format in ['md', 'html', 'pdf']:
        # Always generate markdown version first
        md_output_file = os.path.join(args.output_dir, f"{base_name}--{model_name.replace(':', '-')}.md")
        exec_summary, bullet_summary_str, participants = get_markdown_document(
            exec_summary,
            bullet_summary,
            md_output_file,
            participants
        )
        print(f"Markdown document saved to {md_output_file}")

        # Generate PDF if requested
        if args.format == 'pdf':
            pdf_output_file = os.path.join(args.output_dir, f"{base_name}--{model_name.replace(':', '-')}.pdf")

            convert_md_to_pdf(
                md_output_file,
                pdf_output_file,
                args.title,
                args.date,
                "RealRecap"
            )
            print(f"PDF document saved to {pdf_output_file}")
    else:
        print(f"Unsupported output format: {args.format}")
