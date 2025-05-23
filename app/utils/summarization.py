import os
import re
import json
import requests
import tempfile
import datetime
from typing import Dict, List, Optional
import docx
from flask import current_app


class SummarizationProcessor:
    def __init__(self, model_name='qwen3-summarizer:14b', ollama_host='http://localhost:11434'):
        self.model_name = model_name
        self.ollama_host = ollama_host
        self.address = f"{ollama_host}/api/generate"

        # Define prompts (from original summarize.py)
        self.bullet_prompt = """What follows is a diarized transcript of a meeting between a number of people in the electric utility
industry in the United States. Topics discussed will be related to that or related to project management and software
development activities to support the electric utility. 

Your output should be in proper markdown format with the following:
1. Use ** for bold text (e.g., **important point**)
2. Use ## for topic headings (level 2 headers) - DO NOT use numbered sections or ### headers
3. Use - for bulleted lists (always use a dash followed by a space)
4. Use blank lines between paragraphs
5. Never use numbered lists for bullet points
6. DO NOT include any numbering (1.0.1, 2.1, etc.) in headers or sections
7. DO NOT create subsections under topic headers

You must include the following sections: 
1. For each key topic discussed: a descriptive name as a level 2 header (##), followed by a bulleted list of key points using dashes (-), with action items highlighted in bold.

Do not include anything else except for the above. No extraneous words, HTML formatting, or numbered sections. \n \n"""

        self.exec_summary_prompt = """What follows is a list of key topics discussed in a meeting between a number of people in the 
electric utility industry. The topic will be related to this, or to project management and software development 
activities to support the electric utility.

Your output MUST BE EXACTLY ONE PARAGRAPH that summarizes the meeting.
- Use plain text with markdown formatting (use ** for bold, not HTML tags)
- Write as a single continuous paragraph with no line breaks
- Focus only on the most critical information
- Keep your summary concise (under 150 words)
- Do not include bullet points, numbered lists, or multiple paragraphs
- No extraneous words or HTML formatting at all

Write only the executive summary paragraph and nothing else.\n \n"""

    def load_transcript(self, transcript_file: str) -> tuple:
        """Load transcript from file and extract participants if possible."""
        _, ext = os.path.splitext(transcript_file)
        participants = None

        if ext.lower() == '.txt':
            with open(transcript_file, 'r') as f:
                transcript = f.read()
        elif ext.lower() == '.docx':
            transcript = self._extract_text_from_docx(transcript_file)
            participants = self._extract_participants_from_docx(transcript_file)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

        return transcript, participants

    def _extract_text_from_docx(self, docx_file: str) -> str:
        """Extract text from MS Teams transcript (.docx) file."""
        current_app.logger.info(f"Loading MS Teams transcript from {docx_file}")
        doc = docx.Document(docx_file)
        transcript_text = []
        current_speaker = None

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Check if this is a speaker line
            if ':' in text and len(text.split(':')[0]) < 50:
                parts = text.split(':', 1)
                speaker_info = parts[0].strip()
                content = parts[1].strip() if len(parts) > 1 else ""

                current_speaker = speaker_info
                if content:
                    transcript_text.append(f"{current_speaker}: {content}")
            else:
                # Content from previous speaker
                if current_speaker:
                    transcript_text.append(f"{current_speaker}: {text}")
                else:
                    transcript_text.append(text)

        return '\n'.join(transcript_text)

    def _extract_participants_from_docx(self, docx_file: str) -> List[str]:
        """Extract unique participants from MS Teams transcript."""
        current_app.logger.info(f"Extracting participants from {docx_file}")
        doc = docx.Document(docx_file)
        participants = set()

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            if ':' in text and len(text.split(':')[0]) < 50:
                speaker_info = text.split(':', 1)[0].strip()

                # Clean up speaker name
                speaker_name = speaker_info

                # Remove timestamps and numbers
                if '(' in speaker_name:
                    speaker_name = speaker_name.split('(')[0].strip()

                speaker_name = re.sub(r'\s+\d+\s*$', '', speaker_name)

                # Only add valid names
                if speaker_name and not re.match(
                        r'^(January|February|March|April|May|June|July|August|September|October|November|December)',
                        speaker_name):
                    participants.add(speaker_name)

        return sorted(list(participants))

    def _ollama_request_think_tags(self, prompt: str) -> str:
        """Make request to Ollama API and extract response after think tags."""
        payload = {
            'model': self.model_name,
            'prompt': prompt,
            'stream': False,
        }

        try:
            response = requests.post(self.address, json=payload, timeout=300)  # 5 minute timeout
            response.raise_for_status()

            response_data = response.json()
            full_response = response_data['response']

            # Extract text after </think> tags
            match = re.search(r'</think>(.*)', full_response, re.DOTALL)
            if match:
                return match.group(1).strip()
            else:
                return full_response

        except requests.exceptions.RequestException as e:
            current_app.logger.error(f"Error making Ollama request: {e}")
            raise RuntimeError(f"Failed to get response from Ollama: {e}")

    def get_bullet_summary(self, transcript: str, participants: Optional[List[str]] = None) -> str:
        """Generate bullet point summary from transcript."""
        current_app.logger.info(f"Generating bullet summary with {self.model_name}")

        # Add participants to prompt if available
        local_bullet_prompt = self.bullet_prompt
        if participants:
            participants_str = ", ".join(participants)
            local_bullet_prompt += f"Meeting participants: {participants_str}\n\n"

        return self._ollama_request_think_tags(local_bullet_prompt + transcript)

    def get_exec_summary(self, bullet_summary: str) -> str:
        """Generate executive summary from bullet summary."""
        current_app.logger.info("Generating executive summary")
        return self._ollama_request_think_tags(self.exec_summary_prompt + bullet_summary)

    def create_markdown_document(self, exec_summary: str, bullet_summary: str,
                                 output_dir: str, title: str = None,
                                 participants: Optional[List[str]] = None) -> str:
        """Generate markdown document."""
        output_file = os.path.join(output_dir, f"summary--{self.model_name.replace(':', '-')}.md")

        with open(output_file, 'w') as f:
            # Executive Summary
            f.write("# Executive Summary\n\n")

            # Clean executive summary
            cleaned_exec = exec_summary.strip()
            cleaned_exec = re.sub(r'^#+\s*.*$', '', cleaned_exec, flags=re.MULTILINE)
            cleaned_exec = re.sub(r'^\d+\.[\d\.]*\s*.*$', '', cleaned_exec, flags=re.MULTILINE)
            cleaned_exec = re.sub(r'\n\s*\n', ' ', cleaned_exec)
            cleaned_exec = cleaned_exec.strip()

            f.write(cleaned_exec)
            f.write('\n\n')

            # Participants section
            if participants and len(participants) > 0:
                f.write("## Meeting Participants\n\n")
                # Create 4-column table
                cols = 4
                rows = (len(participants) + cols - 1) // cols

                f.write("| | | | |\n")
                f.write("| --- | --- | --- | --- |\n")

                for i in range(rows):
                    row = "| "
                    for j in range(cols):
                        idx = i + j * rows
                        if idx < len(participants):
                            row += participants[idx] + " | "
                        else:
                            row += " | "
                    f.write(row.strip() + "\n")
                f.write('\n\n')

            # Key Discussion Topics
            f.write("# Key Discussion Topics and Decisions\n\n")

            # Clean bullet summary
            cleaned_summary = bullet_summary
            cleaned_summary = cleaned_summary.replace("<strong>", "**").replace("</strong>", "**")
            cleaned_summary = cleaned_summary.replace("<ul>", "").replace("</ul>", "")
            cleaned_summary = cleaned_summary.replace("<li>", "- ").replace("</li>", "")
            cleaned_summary = cleaned_summary.replace("<p>", "").replace("</p>", "\n\n")

            # Fix numbered sections
            cleaned_summary = re.sub(r'^\d+\.[\d\.]*\s*(.*?)$', r'## \1', cleaned_summary, flags=re.MULTILINE)
            cleaned_summary = cleaned_summary.replace("### ", "## ")

            f.write(cleaned_summary)

        current_app.logger.info(f"Markdown document saved to {output_file}")
        return output_file

    def create_pdf_document(self, markdown_file: str, output_dir: str, title: str = None) -> str:
        """Convert markdown to PDF using the existing md_to_pdf module."""
        from app.utils.pdf_generator import convert_md_to_pdf

        base_name = os.path.splitext(os.path.basename(markdown_file))[0]
        pdf_file = os.path.join(output_dir, f"{base_name}.pdf")

        current_date = datetime.datetime.now().strftime("%B %d, %Y")

        result = convert_md_to_pdf(
            markdown_file,
            pdf_file,
            title=title or "Meeting Summary",
            date=current_date,
            author="RealRecap"
        )

        if result:
            current_app.logger.info(f"PDF document saved to {pdf_file}")
            return pdf_file
        else:
            raise RuntimeError("Failed to generate PDF")

    def process_transcript(self, transcript_file: str, title: str = None,
                           output_format: str = 'pdf') -> Dict:
        """Main processing function."""
        current_app.logger.info(f"Processing transcript: {transcript_file}")

        # Create temporary output directory
        with tempfile.TemporaryDirectory() as temp_dir:
            # Load transcript
            transcript, participants = self.load_transcript(transcript_file)

            # Generate summaries
            bullet_summary = self.get_bullet_summary(transcript, participants)
            exec_summary = self.get_exec_summary(bullet_summary)

            # Create markdown document
            markdown_file = self.create_markdown_document(
                exec_summary, bullet_summary, temp_dir, title, participants
            )

            output_files = [markdown_file]

            # Generate additional formats
            if output_format == 'pdf':
                pdf_file = self.create_pdf_document(markdown_file, temp_dir, title)
                output_files.append(pdf_file)
            elif output_format == 'html':
                # HTML generation would go here
                pass

            return {
                'status': 'completed',
                'executive_summary': exec_summary,
                'bullet_summary': bullet_summary,
                'participants': participants,
                'output_files': output_files
            }